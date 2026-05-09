"""文档加载的内部工具函数。

每个函数处理一种文件类型，返回 (text, warnings)。
对外接口在 document_loader.py 的 load_document()。
"""

from __future__ import annotations
import base64
import shutil
import subprocess
from pathlib import Path


# ============================================================
# 文件大小限制
# ============================================================

MAX_PDF_SIZE_MB = 30
MAX_TEXT_SIZE_MB = 10


def _check_size(path: Path, max_mb: int) -> None:
    """检查文件大小，超过 max_mb 抛 ValueError。"""
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > max_mb:
        raise ValueError(
            f"文件 {path} 大小 {size_mb:.1f} MB，"
            f"超过限制 {max_mb} MB。"
            f"请先压缩或分割文件。"
        )


# ============================================================
# PDF: base64 直传（Tier 1）
# ============================================================

def load_pdf_as_base64(path: Path) -> str:
    """读取 PDF 文件，返回 base64 字符串。
    
    用于支持原生 PDF 的 provider（Anthropic 等）。
    """
    _check_size(path, MAX_PDF_SIZE_MB)
    pdf_bytes = path.read_bytes()
    return base64.standard_b64encode(pdf_bytes).decode("utf-8")


# ============================================================
# PDF: 转图像列表（Tier 2，未实现）
# ============================================================

def load_pdf_as_images(path: Path) -> list[tuple[str, str]]:
    """把 PDF 转成图像列表。
    
    返回：[(media_type, base64_data), ...]
    
    依赖 pdf2image（需要系统装 poppler）。第一版未实现，
    用户走 Tier 2 时给清晰错误信息。
    """
    raise NotImplementedError(
        "PDF 转图像（Tier 2）尚未实现。\n"
        "可选方案：\n"
        "  1. 改用支持原生 PDF 的 provider（Anthropic / OpenAI GPT-4o）\n"
        "  2. 用网页端 LLM 把论文转成 markdown 文本后用 .md 文件输入\n"
        "  3. 等待后续版本支持 pdf2image"
    )


# ============================================================
# PDF: 抽文本（Tier 3，降级路径）
# ============================================================

def load_pdf_as_text(path: Path) -> tuple[str, list[str]]:
    """用 pypdf 抽取 PDF 纯文本。
    
    返回：(text, warnings)
    
    警告信息会提醒用户公式可能丢失。
    """
    _check_size(path, MAX_PDF_SIZE_MB)
    
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ImportError(
            "需要 pypdf 库。请运行：pip install pypdf"
        ) from e

    warnings: list[str] = []
    reader = PdfReader(path)
    
    text_parts = []
    for i, page in enumerate(reader.pages):
        try:
            text_parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            warnings.append(
                f"PDF 第 {i+1} 页文本提取失败：{type(e).__name__}: {e}"
            )

    text = "\n\n".join(text_parts).strip()
    
    if not text:
        warnings.append(
            "从 PDF 中未提取到任何文本。"
            "PDF 可能是扫描件或图像，需要 OCR 处理。"
        )
    else:
        warnings.append(
            "通过 pypdf 提取 PDF 文本，数学公式和图表可能丢失或损坏。"
            "如论文含复杂公式，建议改用支持原生 PDF 的 provider，"
            "或在网页端 LLM 上预处理后再输入。"
        )

    return text, warnings


# ============================================================
# DOCX: 抽文本
# ============================================================

def load_docx_as_text(path: Path) -> tuple[str, list[str]]:
    """优先用 pandoc 转 markdown（公式保留），失败时 fallback 到 python-docx。"""
    
    # 先试 pandoc
    if _has_pandoc():
        try:
            return _load_docx_via_pandoc(path)
        except Exception as e:
            warnings = [f"pandoc 转换失败：{e}，fallback 到 python-docx"]
            text, more_warnings = _load_docx_via_python_docx(path)
            return text, warnings + more_warnings
    
    # 没装 pandoc，直接走 python-docx + 警告
    text, warnings = _load_docx_via_python_docx(path)
    warnings.insert(0, 
        "未检测到 pandoc。安装 pandoc 后 docx 公式可以被正确保留：\n"
        "  conda install -c conda-forge pandoc"
    )
    return text, warnings


def _has_pandoc() -> bool:
    return shutil.which("pandoc") is not None


def _load_docx_via_pandoc(path: Path) -> tuple[str, list[str]]:
    _check_size(path, MAX_TEXT_SIZE_MB)
    result = subprocess.run(
        ["pandoc", str(path), "-t", "markdown", "--wrap=none"],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    if result.returncode != 0:
        raise RuntimeError((result.stderr or result.stdout).strip())

    text = result.stdout.strip()
    warnings: list[str] = []
    if not text:
        warnings.append(f"pandoc extracted no previewable text from {path.name}.")
    return text, warnings


def _load_docx_via_python_docx(path: Path) -> tuple[str, list[str]]:
    _check_size(path, MAX_TEXT_SIZE_MB)
    try:
        from docx import Document
    except ImportError as error:
        raise ImportError("Install python-docx to read .docx files.") from error

    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts).strip()
    warnings = [
        "Loaded docx with python-docx fallback. Equations and complex formatting may be incomplete."
    ]
    if not text:
        warnings.append(f"python-docx extracted no previewable text from {path.name}.")
    return text, warnings


# ============================================================
# 纯文本：txt / md
# ============================================================

def load_plain_text(path: Path) -> tuple[str, list[str]]:
    """读取纯文本文件（txt / md）。"""
    _check_size(path, MAX_TEXT_SIZE_MB)
    
    warnings: list[str] = []
    
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        warnings.append(
            f"文件 {path.name} 不是 UTF-8 编码，使用替换模式读取，"
            f"可能有字符丢失。建议把文件转存为 UTF-8。"
        )
        text = path.read_text(encoding="utf-8", errors="replace")
    
    return text.strip(), warnings
